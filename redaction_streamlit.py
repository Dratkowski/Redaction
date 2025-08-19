# Requirements:
# pip install streamlit python-docx pymupdf spacy
# python -m spacy download en_core_web_sm

import streamlit as st
import os
import re
from collections import defaultdict
from io import BytesIO
from docx import Document
import spacy

# Load spaCy model
nlp = spacy.load("en_core_web_sm")

# Try importing PyMuPDF
try:
    import fitz  # PyMuPDF
    pdf_enabled = True
except ImportError:
    fitz = None
    pdf_enabled = False
    st.warning("PyMuPDF not installed. PDF files will not be supported. Install with `pip install PyMuPDF`.")

# Regex patterns for PII
patterns = {
    'EMAIL': r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[A-Z|a-z]{2,}',
    'PHONE': r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
    'SSN': r'\d{3}-\d{2}-\d{4}',
    'CREDIT_CARD': r'\d{4}[- ]?\d{4}[- ]?\d{4}[- ]?\d{4}',
    'IP_ADDRESS': r'\b(?:\d{1,3}\.){3}\d{1,3}\b',
    'DOB': r'(?:\b(?:0?[1-9]|1[0-2])[-/](?:0?[1-9]|[12]\d|3[01])[-/](?:19|20)\d{2}\b)|(?:\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+(?:0?[1-9]|[12]\d|3[01]),?\s+(?:19|20)\d{2}\b)',
    'ADDRESS': r'\b\d{1,5}\s+[A-Za-z\s]+(?:St|Street|Ave|Avenue|Rd|Road|Blvd|Boulevard|Ln|Lane|Dr|Drive),\s*[A-Za-z\s]+,\s*[A-Z]{2}\s*\d{5}\b',
    'DRIVER_LICENSE': r'(?:\b[A-Z]{1,2}\d{3}-\d{3}-\d{2}-\d{3}-\d{1,2}\b)|(?:\b\d{7,9}\b)'
}

# --- Text extraction ---
def extract_text(file):
    ext = os.path.splitext(file.name)[1].lower()
    if ext == '.pdf':
        if not pdf_enabled:
            st.error("PDF processing is disabled because PyMuPDF is not installed.")
            return None, None
        doc = fitz.open(stream=file.read(), filetype="pdf")
        page_texts = [(page.get_text(), page.number + 1) for page in doc]
        full_text = "\n".join(text for text, _ in page_texts)
        doc.close()
        return full_text, page_texts
    elif ext == '.docx':
        doc = Document(file)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        page_texts = [(full_text, 1)]
        return full_text, page_texts
    else:
        st.error("Unsupported file type")
        return None, None

# --- Context extraction ---
def get_context(text, start, end, pii_value, page_texts):
    page_number = 1
    char_offset = 0
    for page_text, page_num in page_texts:
        if start >= char_offset and start < char_offset + len(page_text):
            page_number = page_num
            break
        char_offset += len(page_text) + 1
    before_words = text[:start].split()[-5:]
    after_words = text[end:].split()[:5]
    context = ' '.join(before_words) + ' ' + pii_value + ' ' + ' '.join(after_words)
    return context, page_number

# --- Regex PII ---
def find_regex_pii(text, page_texts):
    pii = defaultdict(lambda: defaultdict(list))
    for pii_type, pattern in patterns.items():
        for match in re.finditer(pattern, text):
            value = match.group()
            start, end = match.start(), match.end()
            context, page_number = get_context(text, start, end, value, page_texts)
            pii[pii_type][value].append((context, page_number))
    return {k: {vk: list(set(vv)) for vk, vv in v.items()} for k, v in pii.items()}

# --- NER PII ---
def find_ner_pii(text, page_texts):
    doc = nlp(text)
    ner_pii = defaultdict(lambda: defaultdict(list))
    for ent in doc.ents:
        if ent.label_ in ['PERSON', 'GPE']:
            value = ent.text
            start, end = ent.start_char, ent.end_char
            context, page_number = get_context(text, start, end, value, page_texts)
            ner_pii[ent.label_][value].append((context, page_number))
    return {k: {vk: list(set(vv)) for vk, vv in v.items()} for k, v in ner_pii.items()}

# --- DOCX redaction ---
def redact_docx(file, approved):
    doc = Document(file)
    for p in doc.paragraphs:
        for value in approved:
            if value in p.text:
                p.text = p.text.replace(value, "[REDACTED]")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- PDF redaction ---
def redact_pdf(file, approved):
    if not pdf_enabled:
        st.error("PDF redaction disabled (PyMuPDF not installed).")
        return None
    doc = fitz.open(stream=file.read(), filetype="pdf")
    for value in approved:
        for page in doc:
            for inst in page.search_for(value):
                page.add_redact_annot(inst, fill=(0,0,0), text="")
    for page in doc:
        page.apply_redactions()
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    doc.close()
    return buffer

# --- Streamlit interface ---
st.title("PII Redactor with Auto-Select All")

file_types = ["docx"]
if pdf_enabled:
    file_types.append("pdf")

uploaded_file = st.file_uploader("Upload PDF or DOCX", type=file_types)

if uploaded_file:
    text, page_texts = extract_text(uploaded_file)
    if text is not None:
        regex_pii = find_regex_pii(text, page_texts)
        ner_pii = find_ner_pii(text, page_texts)

        # Merge regex and NER
        pii_groups = {}
        all_types = set(list(regex_pii.keys()) + list(ner_pii.keys()))
        for t in all_types:
            pii_groups[t] = {}
            if t in regex_pii:
                pii_groups[t].update(regex_pii[t])
            if t in ner_pii:
                for vk, vv in ner_pii[t].items():
                    if vk in pii_groups[t]:
                        pii_groups[t][vk].extend(vv)
                        pii_groups[t][vk] = list(set(vv))
                    else:
                        pii_groups[t][vk] = vv

        st.header("Select PII to Redact (All occurrences auto-selected)")
        approved = []

        for pii_type, values in pii_groups.items():
            with st.expander(pii_type):
                for value, contexts in values.items():
                    context_text = " | ".join([f"{ctx} (Page {pg})" for ctx, pg in contexts])
                    if st.checkbox(f"{value} → {context_text}", key=value):
                        approved.append(value)

        if st.button("Apply Redactions"):
            ext = os.path.splitext(uploaded_file.name)[1].lower()
            uploaded_file.seek(0)
            if ext == '.pdf':
                redacted_file = redact_pdf(uploaded_file, approved)
                if redacted_file:
                    st.download_button("Download Redacted PDF", redacted_file,
                                       file_name=f"{uploaded_file.name.rsplit('.',1)[0]}_redacted.pdf")
            elif ext == '.docx':
                redacted_file = redact_docx(uploaded_file, approved)
                st.download_button("Download Redacted DOCX", redacted_file,
                                   file_name=f"{uploaded_file.name.rsplit('.',1)[0]}_redacted.docx")
